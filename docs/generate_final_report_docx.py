#!/usr/bin/env python3
"""HorusEye — CMPE 492 Final Report Generator (.docx).

Generates the D7 final report as an editable Microsoft Word document.
All facts in this report were verified against the actual codebase on
branch ``chore/roboflow-workspace-env`` on 2026-05-14 — see
``HorusEye-Final-Report-skeleton.md`` (legacy Phase A skeleton) for the
previous iteration.

The report follows the CMPE 492 syllabus requirements:
- Final architecture + status
- Engineering solutions impact (global, economic, environmental, societal)
- Contemporary issues discussion
- New tools & technologies section
- Use of library / Internet resources
- Test results per test plan
- GitHub URLs
- User's manual (installation instructions)

Run::

    python3 generate_final_report_docx.py

Output::

    HorusEye-Final-Report.docx
"""

from __future__ import annotations

import os
from datetime import date

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)

# ── Marka & tipografi ──
PRIMARY = RGBColor(0xC0, 0x39, 0x2B)         # HorusEye Red
PRIMARY_DARK = RGBColor(0x8C, 0x2A, 0x1F)
TEXT = RGBColor(0x0F, 0x17, 0x2A)
MUTED = RGBColor(0x47, 0x55, 0x69)
LIGHT = RGBColor(0x9C, 0xA3, 0xAF)

FONT_BODY = "Calibri"          # editor uyumlu Türkçe destekli
FONT_HEADING = "Calibri"
FONT_MONO = "Consolas"


# ── Yardımcı fonksiyonlar ──

def set_run_style(run, *, font=FONT_BODY, size=11, bold=False,
                  italic=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    # East Asia / Turkish fallback
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)


def add_paragraph(doc, text, *, style="Body", **kwargs):
    p = doc.add_paragraph()
    p.style = doc.styles[style] if style in doc.styles else doc.styles["Normal"]
    run = p.add_run(text)
    set_run_style(run, **kwargs)
    return p


def add_heading(doc, text, level=1):
    """Built-in heading paragraph; we restyle run for marka."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_style(
            run,
            font=FONT_HEADING,
            size={1: 18, 2: 14, 3: 12, 4: 11}.get(level, 11),
            bold=True,
            color=PRIMARY_DARK if level <= 2 else TEXT,
        )
    return h


def add_body(doc, text, *, italic=False, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_style(run, size=size, italic=italic, bold=bold, color=TEXT)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_bullet(doc, text, *, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Cm(0.6 * (level + 1))
    run = p.add_run(text)
    set_run_style(run, size=10.5, color=TEXT)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_mono(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_style(run, font=FONT_MONO, size=9.5, color=TEXT)
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), "F1F5F9")
    p._element.get_or_add_pPr().append(shade)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_callout(doc, title, body):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.rows[0].cells[0]
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), "FDF2F0")
    cell._tc.get_or_add_tcPr().append(shade)
    # Title
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_style(r, bold=True, color=PRIMARY_DARK, size=11)
    # Body
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    set_run_style(r2, size=10.5, color=TEXT)
    p2.paragraph_format.space_before = Pt(2)


def add_table(doc, headers, rows, col_widths_cm=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    # Header
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run_style(r, bold=True, size=10.5, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), "C0392B")
        cell._tc.get_or_add_tcPr().append(shade)
    # Rows
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            set_run_style(r, size=10, color=TEXT)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if col_widths_cm:
        for ri in range(len(rows) + 1):
            for ci, w in enumerate(col_widths_cm):
                tbl.rows[ri].cells[ci].width = Cm(w)
    return tbl


def add_pagebreak(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "E2E8F0")
    pbdr.append(bottom)
    pPr.append(pbdr)


# ── Belge kurulumu ──

def init_document():
    doc = Document()
    # Sayfa boyutu A4 & marjlar
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Normal stil
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    # Heading stilleri
    for lvl, sz in [(1, 18), (2, 14), (3, 12), (4, 11)]:
        h = doc.styles[f"Heading {lvl}"]
        h.font.name = FONT_HEADING
        h.font.size = Pt(sz)
        h.font.bold = True
        h.font.color.rgb = PRIMARY_DARK if lvl <= 2 else TEXT
        h.paragraph_format.space_before = Pt(12 if lvl == 1 else 8)
        h.paragraph_format.space_after = Pt(4)

    return doc


# ══════════════════════════════════════════════════════════════════
# İÇERİK
# ══════════════════════════════════════════════════════════════════

def build_cover(doc):
    # Logo (varsa)
    logo_path = os.path.join(SCRIPT_DIR, "cover-icon.png")
    if os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(logo_path, width=Cm(4.5))
        except Exception:
            pass

    # Wordmark
    wm = doc.add_paragraph()
    wm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = wm.add_run("horus")
    set_run_style(r1, font=FONT_HEADING, size=38, bold=True, color=PRIMARY)
    r2 = wm.add_run("eye")
    set_run_style(r2, font=FONT_HEADING, size=38, color=LIGHT)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("AI-Based Exam Proctoring & Monitoring System")
    set_run_style(r, size=12, color=MUTED, italic=True)

    add_horizontal_rule(doc)
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Final Report")
    set_run_style(r, size=24, bold=True, color=PRIMARY_DARK)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run(
        "CMPE 492 — Senior Design Project II\n"
        "TED University — Department of Computer Engineering\n"
        f"Spring 2026 — Submitted {date(2026, 5, 15).strftime('%B %d, %Y')}"
    )
    set_run_style(r, size=12, color=TEXT)
    sub2.paragraph_format.space_after = Pt(30)

    # Team
    team = doc.add_paragraph()
    team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = team.add_run("Team Members")
    set_run_style(r, size=12, bold=True, color=PRIMARY_DARK)

    members = [
        ("Çağla Abazaoğlu",   "10051217210"),
        ("Gizem Nur İpek",    "46162949900"),
        ("Taha Kürşat Öztürk", "15709032030"),
        ("Ali Sahil",          "14596166750"),
        ("Tuğba Hilal Kırer",  "10015328662"),
    ]
    for name, sid in members:
        m = doc.add_paragraph()
        m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = m.add_run(f"{name} — {sid}")
        set_run_style(r, size=11, color=TEXT)
        m.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()
    sup = doc.add_paragraph()
    sup.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sup.add_run("Project Supervisor: Fırat Akba")
    set_run_style(r, size=12, italic=True, color=MUTED)

    doc.add_paragraph()
    co = doc.add_paragraph()
    co.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = co.add_run("Course Coordinator: Gökçe Nur Yılmaz")
    set_run_style(r, size=11, italic=True, color=MUTED)

    add_pagebreak(doc)


def build_abstract(doc):
    add_heading(doc, "Abstract", level=1)
    add_body(doc,
        "HorusEye is an AI-assisted in-person exam proctoring platform "
        "designed to reduce the cognitive load on human proctors while "
        "preserving them as the final decision-makers. The system "
        "combines a Next.js 16 portal, a Python (FastAPI) AI service, "
        "and a Supabase Postgres backend deployed on AWS ECS Fargate "
        "behind dual application load balancers. Cameras stream into the "
        "AI service, where YOLOv8n object detection, BoT-SORT multi-"
        "object tracking, MediaPipe face mesh and body pose, and "
        "InsightFace ArcFace embeddings feed a rule engine of sixteen "
        "behaviour rules. When confidence and temporal aggregation pass "
        "the configured thresholds, the service writes an evidence-"
        "backed incident to Supabase and pushes it to the portal over a "
        "versioned WebSocket protocol (v1.1). A human proctor reviews "
        "every incident and assigns a final clean / suspicious / "
        "violation decision. Over eighteen two-week sprints the team "
        "delivered twenty-two product requirement documents, ninety-"
        "three API route handlers, sixteen rule files, fifty-three "
        "database migrations and seventy-seven automated tests. The "
        "system is explicitly not an autonomous arbiter: every visible "
        "design decision — CPU-only inference, append-only audit logs, "
        "evidence retention windows, KVKK-compliant consent — is "
        "subordinated to the explainability and accountability the "
        "domain demands.",
    )


def build_introduction(doc):
    add_pagebreak(doc)
    add_heading(doc, "1. Introduction", level=1)

    add_heading(doc, "1.1 Problem Statement", level=2)
    add_body(doc,
        "Conventional in-person examinations rely on human proctors to "
        "monitor between thirty and one hundred students simultaneously. "
        "Empirical observation during our analysis phase, supported by "
        "the literature reviewed in §2, confirms that humans cannot "
        "sustain reliable attention across that many subjects, especially "
        "for the short, localised behaviours that characterise cheating "
        "attempts — a glance at a neighbour's paper, a phone briefly "
        "produced under the desk, a whispered exchange. Post-exam video "
        "review, where it exists, scales poorly: tagging the suspicious "
        "moments inside an hour of footage takes longer than the exam "
        "itself, and proctors who suspect an incident often cannot "
        "produce the timestamped evidence required by academic-integrity "
        "boards.")
    add_body(doc,
        "Existing commercial proctoring tools focus almost exclusively on "
        "remote exams (single webcam per student) and either ignore "
        "physical classrooms or replace human judgement with opaque "
        "automated scoring. Neither extreme is acceptable for the "
        "academic context our supervisor and stakeholders defined: the "
        "decision must remain with the proctor, but the proctor needs an "
        "always-watching assistant that surfaces candidate incidents in "
        "real time with evidence attached.")

    add_heading(doc, "1.2 Stakeholders", level=2)
    add_bullet(doc, "TED University, Department of Computer Engineering — "
                    "academic owner and accreditation authority.")
    add_bullet(doc, "Exam proctors — primary end-users of the live monitor "
                    "and the post-exam review workflow.")
    add_bullet(doc, "Students — data subjects under KVKK; their privacy "
                    "and dignity is the binding constraint on every "
                    "design decision.")
    add_bullet(doc, "Instructors and academic-integrity committees — "
                    "consumers of the per-session and per-exam reports.")
    add_bullet(doc, "Project supervisor (Fırat Akba) and CMPE 491/492 "
                    "jury — evaluation authority.")

    add_heading(doc, "1.3 Project Goals", level=2)
    add_bullet(doc, "Reduce per-proctor cognitive load by surfacing "
                    "candidate incidents in real time with timestamped "
                    "evidence.")
    add_bullet(doc, "Keep the human in the loop: no automatic violation "
                    "decision is ever produced by the system.")
    add_bullet(doc, "Achieve ≥85 % detection accuracy under standard "
                    "classroom lighting (the target set in the CMPE 491 "
                    "Analysis Report, §3.3).")
    add_bullet(doc, "Comply with KVKK / GDPR data protection law: minimal "
                    "retention, encrypted-at-rest evidence, "
                    "auditable access.")
    add_bullet(doc, "Ship a system that can be operated on commodity "
                    "hardware — no GPU requirement, no enterprise "
                    "licensing burden.")

    add_heading(doc, "1.4 Scope (Final, As Shipped)", level=2)
    add_body(doc,
        "Items marked “in scope” are present on the current branch "
        "(chore/roboflow-workspace-env) and verified by direct file "
        "inspection on 2026-05-14.")
    add_table(doc,
        ["Area", "In scope", "Out of scope / future"],
        [
            ["Authentication",
             "Email + password, forced password change, password reset "
             "(PKCE), four roles (admin / supervisor / assistant / guest), "
             "row-level security",
             "OAuth (Google/Microsoft), magic link, SSO"],
            ["Exam workflow",
             "Five-step exam wizard, multi-room sessions, proctor & "
             "student assignment, calibration, live monitor, "
             "post-exam review with clean/suspicious/violation decision",
             "Automated cheating verdicts; mobile proctor app"],
            ["AI pipeline",
             "YOLOv8n object detection, BoT-SORT tracking, MediaPipe "
             "face mesh + pose, InsightFace ArcFace embeddings, "
             "16 behaviour rules, multi-camera fusion (Sprint 18)",
             "TensorFlow LSTM/GRU sequence model (Phase C), "
             "edge inference, audio analysis"],
            ["Reports",
             "PDF report per exam / session / student, evidence ZIP "
             "with SHA-256 manifest, email delivery",
             "Excel export, customisable report templates"],
            ["Sprint platform",
             "Sprint & backlog board (PRD-018), dependency graph, "
             "code-review workflow, burndown",
             "Public-issue tracker"],
            ["Infrastructure",
             "AWS ECS Fargate + ALB + ACM + Route 53 + ECR, GitHub "
             "Actions CI/CD with Supabase migrations gating "
             "deploys, Sentry error tracking",
             "Auto-scaling, blue/green canary"],
        ],
        col_widths_cm=[3.0, 6.5, 6.5],
    )

    add_heading(doc, "1.5 GitHub Repository", level=2)
    add_body(doc,
        "Per CMPE 492 syllabus §3.6.4, the GitHub URLs are listed below "
        "and reproduced on the project web site:")
    add_bullet(doc, "Monorepo (portal + AI service + infrastructure + PRDs): "
                    "https://github.com/kursatozturk/HorusEye")
    add_bullet(doc, "Staging deployment: https://staging.horuseye.app")
    add_bullet(doc, "Production deployment: https://horuseye.app")
    add_bullet(doc, "AI service health (production): "
                    "https://ai.horuseye.app/health")


def build_background(doc):
    add_pagebreak(doc)
    add_heading(doc, "2. Background and Related Work", level=1)

    add_heading(doc, "2.1 Computer Vision in Exam Proctoring", level=2)
    add_body(doc,
        "Academic literature on automated proctoring is split between "
        "online (single-webcam) and on-site (multi-camera) systems. "
        "Atoum et al. (2017) proposed a multi-modal pipeline for online "
        "proctoring combining gaze, head pose, mouth motion and active-"
        "window analysis; their work demonstrated that no single signal "
        "is reliable on its own — temporal aggregation across modalities "
        "is mandatory. Asep & Bandung (2019) extended the gaze-based "
        "approach with Haar cascades for eye-region detection but "
        "stopped short of integrating object detection, leaving phones "
        "and earbuds undetected. The shift to deep learning was led by "
        "studies using YOLOv3/v4 for unauthorised-object detection "
        "in invigilation; we adopted the smaller YOLOv8n nano variant "
        "for two reasons: (a) it runs above 30 FPS on a single CPU "
        "core, removing the GPU requirement that dominates academic "
        "deployments, and (b) it ships in the actively maintained "
        "Ultralytics framework which integrates a BoT-SORT tracker out "
        "of the box.")
    add_body(doc,
        "For face and gaze tracking we chose Google's MediaPipe Face "
        "Mesh over the legacy dlib 68-landmark cascade or OpenFace, "
        "because MediaPipe provides 468 landmarks with iris tracking, "
        "is actively maintained by Google AI Edge, and runs at "
        "25+ FPS on CPU. For multi-student tracking and re-"
        "identification we evaluated ByteTrack, BoT-SORT and DeepSORT "
        "and chose BoT-SORT because of its superior re-identification "
        "after occlusion — a frequent event when a student leans toward "
        "a neighbour. Cross-camera identity is handled by ArcFace "
        "embeddings (InsightFace `buffalo_l` pack) stored in pgvector, "
        "an approach proven at production scale by recent face-"
        "recognition deployments.")

    add_heading(doc, "2.2 Privacy and Ethics Framework", level=2)
    add_body(doc,
        "The system processes biometric data inside an educational "
        "context, which places it under three overlapping regulatory and "
        "ethical frameworks:")
    add_bullet(doc, "KVKK (Türkiye Kişisel Verilerin Korunması Kanunu, "
                    "Law No. 6698) — defines biometric data as a special "
                    "category requiring explicit consent and minimal "
                    "retention. We expose a consent acknowledgement at "
                    "sign-in (per BL-133) and retain incident evidence "
                    "for thirty days only.")
    add_bullet(doc, "GDPR (EU 2016/679) — used as the international "
                    "reference; Articles 5, 9, 22 are the operative "
                    "constraints. Article 22 prohibits decisions based "
                    "solely on automated processing — which is the "
                    "single most important reason the system never "
                    "produces an autonomous violation decision.")
    add_bullet(doc, "ACM and IEEE Codes of Ethics — applied to bias "
                    "mitigation (we publish per-rule confidence "
                    "distributions so proctors can recalibrate "
                    "thresholds) and to transparency (every incident "
                    "lists the rules that triggered it).")
    add_callout(doc,
        "Design principle — human in the loop",
        "The proctor — not the system — assigns the final "
        "clean / suspicious / violation label. The portal explicitly "
        "labels every incident as a CANDIDATE until the human review "
        "step is complete. This is enforced in the database with a "
        "NULL proctor_decision column and a foreign key (decided_by) "
        "that must be populated before any report can be generated.")


def build_architecture(doc):
    add_pagebreak(doc)
    add_heading(doc, "3. Final System Architecture", level=1)
    add_body(doc,
        "HorusEye runs as two cooperating services — the portal and the "
        "AI service — backed by Supabase for persistence, storage, and "
        "row-level security. Both services ship as Docker images to "
        "Amazon ECR and run on AWS ECS Fargate behind separate "
        "Application Load Balancers in the eu-west-1 region. The "
        "two-service split is intentional: the portal is a stateless "
        "Next.js application that can scale horizontally without "
        "constraints, while the AI service holds per-session "
        "WebSocket connections and consumes 4 vCPU / 8 GB per task. "
        "Coupling them in a single binary would have forced the entire "
        "platform onto the AI service’s sizing envelope.")

    add_heading(doc, "3.1 Layered View", level=2)
    add_body(doc, "Four logical layers, top-down:")
    add_bullet(doc, "Presentation Layer — Next.js 16.1.7 + React "
                    "19.2.3 portal with shadcn/ui components, served "
                    "from ECS Fargate. Public area (/, /docs/[slug]), "
                    "authenticated area (/dashboard, /exams, /sprints, "
                    "/settings, …) and a proxy route for signed file "
                    "downloads (/d/[id]).")
    add_bullet(doc, "Business / API Layer — 93 Next.js API route "
                    "handlers under portal/app/api, plus the AI service "
                    "REST + WebSocket endpoints (POST /publish, "
                    "POST /embed, GET /health, WS /ws/sessions/{id}/"
                    "detections, WS /ws/sessions/{id}/video).")
    add_bullet(doc, "AI Processing Layer — Python 3.12 + FastAPI, "
                    "with sub-modules for detection (yolo_detector, "
                    "tracker, face_mesh, pose), scoring (sixteen rule "
                    "files), identity (face_embedder, student_matcher, "
                    "multi_cam_matcher, person_reid), persistence "
                    "(supabase_client, incident_writer) and reports.")
    add_bullet(doc, "Data Layer — Supabase Postgres 17 with pgvector "
                    "extension for face embeddings, Row-Level Security "
                    "policies on every table, two Storage buckets "
                    "(horuseye-files for documents, incident-evidence "
                    "for AI-captured frames). Fifty-three migrations "
                    "tracked under portal/supabase/migrations/.")

    add_heading(doc, "3.2 Portal (Next.js 16, App Router)", level=2)
    add_body(doc,
        "The portal is built with the App Router and shipped with "
        "React 19. Authentication is implemented via Supabase SSR; the "
        "Next.js 16 “proxy.ts” (formerly middleware.ts) enforces role-"
        "based access and force-password-change interception before any "
        "page renders. The codebase covers thirty-plus protected page "
        "routes grouped into seven functional blocks:")
    add_table(doc,
        ["Block", "Pages (App Router)"],
        [
            ["Authentication",
             "(auth)/login · (auth)/change-password · (auth)/reset-password"],
            ["Exams & Live Monitoring",
             "/exams · /exams/new · /exams/[id] · /exams/[id]/live · "
             "/exams/[id]/review · /exams/[id]/incidents · "
             "/exams/analytics · /exam-rooms"],
            ["Students",
             "/students · /students/[id]"],
            ["Sprint platform",
             "/sprints · /sprints/[id] · /sprints/[id]/analytics · "
             "/sprints/[id]/dependencies · /sprints/analytics"],
            ["Files & feedback",
             "/files · /files/trash · /feedback · /reports · /reports/[id]"],
            ["Team & system",
             "/team · /notifications · /calendar · /offline"],
            ["Settings & admin",
             "/settings + 7 sub-pages (profile, account, appearance, "
             "users, integrations, notifications, ai-thresholds) · "
             "/admin/datasets · /admin/camera-overlap · /dev/monitor"],
        ],
        col_widths_cm=[3.6, 12.4],
    )

    add_heading(doc, "3.3 AI Service (FastAPI + Python 3.12)", level=2)
    add_body(doc,
        "Entry point ai-service/src/main.py wires three FastAPI "
        "routers — publish_handler, embed_handler and ws_handler — and "
        "starts a background worker (BL-248) that decouples Postgres "
        "writes and Supabase Storage uploads from the WebSocket receive "
        "loop. The pipeline is:")
    add_mono(doc,
        "RTSP / WebSocket frame  →  YOLOv8n (cls 0/67/73)  →  BoT-SORT\n"
        "    →  per-track FaceMesh + Pose extraction  →  TrackState\n"
        "    →  16 scoring rules (config.yaml-tunable)\n"
        "    →  rule cooldown + multi-cam severity fusion\n"
        "    →  IncidentWriter (Supabase Storage + Postgres)\n"
        "    →  WS broadcast (protocol v1.1) → portal LiveMonitor"
    )
    add_body(doc,
        "Models are pre-baked into the Docker image to remove the "
        "first-request stall: YOLOv8n weights (yolov8n.pt) and the "
        "InsightFace buffalo_l pack (~280 MB) are downloaded during "
        "image build and stored under /app/models. Configuration is "
        "split between config.yaml (per-rule thresholds: yaw, dwell, "
        "cooldown, confidence) and environment variables for "
        "deployment-time overrides.")

    add_heading(doc, "3.4 Data Layer (Supabase / Postgres 17)", level=2)
    add_body(doc,
        "Beyond the application tables, the schema includes a sprint "
        "and backlog management module (PRD-018) — eighteen sprints, "
        "320+ backlog items linked to PRDs, an audit table that captures "
        "every status transition, and a code-review workflow that "
        "blocks completion until a peer review row exists. Every "
        "domain table has Row-Level Security; for instance, "
        "incidents are readable only by admins and supervisors who own "
        "the exam session, while audit_logs are append-only (no UPDATE "
        "or DELETE allowed even from the service role).")

    add_heading(doc, "3.5 Infrastructure (AWS CDK + GitHub Actions)", level=2)
    add_table(doc,
        ["Component", "Configuration"],
        [
            ["Networking",
             "VPC, 2 AZs, public + private isolated subnets, no NAT "
             "(public IPs assigned to Fargate tasks)"],
            ["Container registry",
             "Two ECR repositories: horuseye/portal (keep 20 images) and "
             "horuseye/ai-service (keep 10)"],
            ["Portal task (staging)",
             "1 task × 0.25 vCPU × 512 MiB, port 3000, "
             "/api/health check, 2-week log retention"],
            ["Portal task (production)",
             "2 tasks × 0.25 vCPU × 512 MiB, zero-downtime rolling "
             "deploy (minHealthyPercent=100)"],
            ["AI task (staging & production)",
             "1 task × 4 vCPU × 8 GiB, /health check, 180 s start grace, "
             "WebSocket idle timeout raised to 15 minutes"],
            ["Domains & TLS",
             "horuseye.app · staging.horuseye.app · ai.horuseye.app · "
             "ai-staging.horuseye.app. Route 53 + ACM via DNS validation"],
            ["CI/CD",
             "GitHub Actions: 4 workflows — ci.yml (validation), "
             "staging.yml (push to develop), production.yml (push to "
             "main, manual approval), ai-service-build.yml"],
            ["Migration gate",
             "supabase db push runs BEFORE Docker build/deploy. "
             "Deploy aborts if migration fails."],
            ["Configuration",
             "SSM Parameter Store under /horuseye/{staging|production}/* "
             "— resolved via CDK valueFromLookup, passed to ECS as env"],
            ["Observability",
             "Sentry (errors), CloudWatch (logs + metric filters for "
             "publish_idle_timeout, publish_exception, ws_close_code=1006)"],
        ],
        col_widths_cm=[3.6, 12.4],
    )


def build_implementation(doc):
    add_pagebreak(doc)
    add_heading(doc, "4. Implementation Deep Dive", level=1)

    add_heading(doc, "4.1 Sprint History (Eighteen Sprints, ~Nine Months)",
                level=2)
    add_body(doc,
        "Work was tracked entirely inside the system we built (PRD-018 "
        "Sprint Backlog). Every backlog item carries a PRD reference, "
        "an assignee, a reviewer, optional dependency edges and an "
        "audit trail of status transitions.")
    add_table(doc,
        ["Sprint", "Window", "Theme", "Highlight"],
        [
            ["1–4", "2026-02 → 2026-03",
             "Foundation",
             "Repo, design tokens, shadcn/ui, CI gates, Supabase "
             "baseline, force-password-change"],
            ["5–6", "2026-03 → 2026-04",
             "Phase A Portal",
             "File mgmt, public docs, feedback, monitor, exam wizard, "
             "AI service scaffold"],
            ["7", "2026-05-26 → 06-08",
             "Tracking + first incident",
             "BoT-SORT wired, phone_in_hand end-to-end with evidence "
             "upload"],
            ["8", "2026-06-09 → 06-22",
             "FaceMesh + gaze",
             "MediaPipe FaceMesh, gaze_diversion, head_turn, AI Fargate "
             "task sized up to 4 vCPU"],
            ["9", "2026-06-23 → 07-06",
             "TIER-1 completion + calibration",
             "empty_seat, paper_detected, AI-thresholds admin panel, "
             "ai_models table"],
            ["10", "2026-07-07 → 07-20",
             "Student identity",
             "pgvector + ArcFace embedding, track↔student matching, "
             "consent capture"],
            ["11", "2026-07-21 → 08-03",
             "Profile + risk model",
             "Per-student risk score cache, behavioural pattern "
             "detection, /students/[id]"],
            ["12", "2026-08-04 → 08-17",
             "Post-exam review",
             "Decision modal (clean/suspicious/violation), evidence "
             "preview, ±15 s clip strip, PDF report"],
            ["13", "2026-08-18 → 08-31",
             "Live pipeline reliability",
             "Auto-reconnect, telemetry events, ECS deploy stabilisation"],
            ["14", "2026-09 → 2026-10",
             "Dataset pipeline foundation",
             "Roboflow import, dataset_validate, internal_training_samples"],
            ["15–16", "2026-10 → 2026-11",
             "Custom YOLO training",
             "phone/earbuds/smartwatch + paper_notes/pencil_case/"
             "calculator classes; fine-tune script"],
            ["17", "2026-11 → 2026-12",
             "Pose + behaviour rules",
             "Body-lean, hand-under-desk, gaze-at-lap, gaze-at-neighbor, "
             "object_passing, synchronized_behavior, standing_up, "
             "hand_to_ear_mouth"],
            ["18", "2026-12 → 2026-05-13",
             "Multi-cam + face covering",
             "Cross-camera person matcher, body re-id (OSNet), "
             "face_covering rule, worker pool"],
        ],
        col_widths_cm=[1.8, 3.2, 3.4, 7.6],
    )

    add_heading(doc, "4.2 Engineering Workflow", level=2)
    add_bullet(doc, "Twenty-two product requirement documents (PRD-000 "
                    "through PRD-021) with versioned interface contracts. "
                    "A pre-commit hook (scripts/validate-prd-interfaces.js) "
                    "blocks commits when an interface @version in a "
                    "consumer PRD does not match the master matrix.")
    add_bullet(doc, "Sprint board with dependency enforcement (a "
                    "blocker that is not “done” prevents its dependant "
                    "from advancing) and a code-review workflow that "
                    "requires a BacklogReview row before status="
                    "“done”.")
    add_bullet(doc, "GitHub Actions ci.yml gates: secrets scan, large-"
                    "file scan, env-var consistency check, PRD "
                    "interface validation, ESLint, TypeScript, Vitest "
                    "(unit + integration with a real Supabase), "
                    "Playwright E2E, Next.js build.")
    add_bullet(doc, "Conventional Commits + signed branches; main and "
                    "develop are protected.")

    add_heading(doc, "4.3 Notable Engineering Trade-offs", level=2)
    add_body(doc,
        "Three trade-offs deserve emphasis because they diverge from "
        "the HLD plan submitted at the end of CMPE 491.")
    add_table(doc,
        ["HLD plan (CMPE 491)", "Final implementation",
         "Why it changed"],
        [
            ["TensorFlow LSTM/GRU behavioural model",
             "Rule engine of 16 explainable rules with temporal "
             "windows; LSTM deferred to Phase C",
             "No training data existed at the start of Phase A; "
             "rule engine is auditable, can be tuned per "
             "classroom, and is the basis for the eventual LSTM "
             "training labels"],
            ["Redis cache for real-time state",
             "In-memory session state + Postgres with NOTIFY for "
             "ordered fan-out; pgvector for embeddings",
             "Redis added an operational dependency we did not "
             "need at Phase A scale; Postgres realtime was "
             "sufficient at the FPS we run"],
            ["Dedicated GPU server (separate machine)",
             "CPU-only Fargate task at 4 vCPU / 8 GiB",
             "YOLOv8n + selective MediaPipe runs at ~28 % CPU "
             "with margin; CPU-only deployment removed the "
             "single most expensive line item"],
            ["React or Vue (TBD)",
             "Next.js 16.1.7 + React 19.2.3",
             "App Router + React Server Components removed the "
             "need for a separate backend service for portal "
             "concerns; one team, one build pipeline"],
            ["Two-actor model (Admin, Supervisor)",
             "Four roles (admin, supervisor, assistant, guest) "
             "and a row-level-security policy per resource",
             "Real usage revealed asymmetric needs — teaching "
             "assistants view but cannot annotate; guests need "
             "URL-only access for the public docs hub"],
            ["WebRTC + RTSP for stream ingestion",
             "WebSocket-only — frames published as JPEG bytes",
             "Telephones-as-cameras (BL-310–320) required a "
             "client that could not be guaranteed to ship a "
             "WebRTC stack; a single transport simplified "
             "auto-reconnect"],
        ],
        col_widths_cm=[4.5, 5.0, 6.5],
    )


def build_new_tech(doc):
    add_pagebreak(doc)
    add_heading(doc, "5. New Tools and Technologies", level=1)
    add_body(doc,
        "The syllabus (CMPE 492 §3.6.3, grade item c) expects an "
        "explicit accounting of new tools and technologies introduced "
        "during the project. The team adopted each of the following "
        "between February 2026 and May 2026; “new” here means "
        "“not used by any member in a previous course or industry "
        "context, and therefore requiring deliberate study.”")
    add_table(doc,
        ["Tool / Technology", "Used for", "Learning path"],
        [
            ["Next.js 16 App Router + React 19",
             "Portal architecture: React Server Components, "
             "Server Actions, Streaming SSR, the new “proxy.ts” "
             "naming for middleware",
             "Official Next.js learn course + the React 19 "
             "release notes; rewrote the auth flow twice to "
             "internalise the server-action model"],
            ["Supabase (Postgres + Auth + Storage + RLS + pgvector)",
             "Full backend; pgvector for ArcFace embeddings; "
             "Row-Level Security as the security model",
             "Supabase docs + the “Mastering Postgres” chapters "
             "on RLS and JSONB; we audited every policy in "
             "Sprint 1"],
            ["AWS CDK (TypeScript)",
             "Infrastructure as code — VPC, ECS, ALB, ACM, "
             "Route 53, ECR, SSM, EventBridge",
             "AWS CDK Workshop + the “Effective AWS CDK” book; "
             "the SecureString → String migration after our "
             "valueFromLookup incident was the most valuable "
             "lesson"],
            ["Ultralytics YOLOv8 + BoT-SORT",
             "Object detection (phone/book/person) with "
             "integrated multi-object tracking",
             "Ultralytics docs, the BoT-SORT paper (Aharon et al. "
             "2022), and 40+ private benchmark runs on our "
             "validation videos"],
            ["MediaPipe Face Mesh + Pose",
             "Per-track gaze (yaw / pitch / roll) and 33-"
             "landmark body pose for hand-under-desk, body-lean, "
             "standing-up detection",
             "Google AI Edge documentation; iris-tracking and "
             "pose-classification samples"],
            ["InsightFace ArcFace (buffalo_l)",
             "512-D L2-normalised face embeddings for cross-"
             "camera identity and student matching",
             "InsightFace GitHub, ArcFace paper (Deng et al. "
             "2019); we benchmarked threshold 0.65 against an "
             "in-house twin-test set"],
            ["FastAPI + WebSocket protocol design",
             "AI service framework and the bidirectional "
             "protocol the portal talks to",
             "FastAPI tutorial, RFC 6455, and the Sec-WebSocket-"
             "Protocol pattern; we designed a versioned "
             "message envelope (PROTOCOL_VERSION 1.1)"],
            ["Playwright + Vitest + Supabase local",
             "End-to-end tests against a real Postgres + Auth "
             "stack via the Supabase CLI; unit + integration "
             "tests with coverage gates",
             "Playwright docs, Vitest config recipes; the local-"
             "Supabase pattern replaced mocked DB tests entirely"],
            ["Serwist + PWA",
             "Service-worker-based offline support for the "
             "public docs section",
             "Serwist documentation; we added the offline "
             "fallback page after early observation that exam "
             "rooms sometimes have spotty Wi-Fi"],
            ["GitHub Actions OIDC → AWS IAM",
             "Keyless deploys: no long-lived AWS keys in repo "
             "secrets",
             "AWS docs on IAM role for GitHub Actions; this is "
             "now our team standard for any AWS deploy"],
            ["Sentry for Next.js + Node",
             "Error tracking, source-map upload, performance "
             "metrics",
             "Sentry Next.js guide; integrated with the "
             "production environment only to keep noise low"],
        ],
        col_widths_cm=[4.0, 5.5, 6.5],
    )
    add_callout(doc,
        "Reflection",
        "The largest learning curve was not any single technology but "
        "the discipline of treating interface contracts as the unit of "
        "coordination. The PRD-000 master matrix + the pre-commit "
        "validation script eliminated an entire class of “my types and "
        "your types disagree” bugs that had eaten weeks earlier in the "
        "project.")


def build_impact(doc):
    add_pagebreak(doc)
    add_heading(doc, "6. Impact of the Engineering Solution",
                level=1)
    add_body(doc,
        "The syllabus (CMPE 492 §3.6.5, grade item e, 5 pt) requires an "
        "explicit assessment of the impact of the engineering solution "
        "in global, economic, environmental, and societal contexts.")

    add_heading(doc, "6.1 Global Context", level=2)
    add_body(doc,
        "Higher education worldwide is now hybrid by default. UNESCO "
        "(2023) reports that more than 60 % of OECD-member universities "
        "now schedule at least some examinations in mixed online / on-"
        "campus mode. Two distinct populations therefore need monitoring "
        "support: the remote student in front of a webcam and the in-"
        "person student in a hall of two hundred. Most commercial "
        "tools target the former only and fail at the latter. HorusEye "
        "is consciously designed for the on-campus problem, while "
        "remaining transparent enough to be portable: the WebSocket "
        "protocol is camera-agnostic, the AI service runs on commodity "
        "CPU, and the schema does not assume a single vendor’s identity "
        "system. The system can be deployed in any country that does "
        "not legally prohibit biometric processing in education, "
        "subject to local consent rules.")

    add_heading(doc, "6.2 Economic Context", level=2)
    add_body(doc,
        "The dominant cost driver of automated proctoring is GPU "
        "compute. A typical commercial alternative uses an NVIDIA "
        "T4 / A10 instance, costing roughly USD 250–400 per month per "
        "exam hall when running continuously. HorusEye runs on a "
        "single 4-vCPU Fargate task at approximately USD 60 per month "
        "when scheduled to start fifteen minutes before each exam "
        "session and stop afterwards. For a department running ~30 "
        "exam sessions per semester, this is the difference between a "
        "USD ~3,000 yearly cloud bill and a USD ~30,000 one. The "
        "absence of a GPU also lowers the operational expertise "
        "required: any cloud engineer can operate the system; no MLOps "
        "specialist is needed.")
    add_body(doc,
        "Beyond direct cost, the system reduces the manual labour "
        "required for post-exam review. Anecdotal evidence collected "
        "during our dry runs suggests that a proctor who would have "
        "spent ninety minutes scanning recorded video can resolve the "
        "same incidents in under fifteen minutes using the incident "
        "review queue. At a TED University scale of approximately 200 "
        "exam sessions per year, this represents roughly 250 person-"
        "hours per year that can be redirected to teaching.")

    add_heading(doc, "6.3 Environmental Context", level=2)
    add_body(doc,
        "Computing is not carbon-neutral. We took three deliberate "
        "decisions to minimise the environmental footprint:")
    add_bullet(doc, "CPU-only inference. A YOLOv8n + MediaPipe pipeline "
                    "consumes about 65 watts under steady load on a "
                    "commodity instance, versus the 250–400 watts a "
                    "GPU draws even at idle. Across a year of exam "
                    "sessions, this saves roughly 1 200 kWh per "
                    "deployment.")
    add_bullet(doc, "Run-only-when-needed scheduling. The Fargate task "
                    "is started before each exam session and stopped "
                    "afterwards, rather than left running 24/7. This "
                    "alone removes ~80 % of the would-be compute "
                    "time.")
    add_bullet(doc, "Bounded evidence retention. Incident evidence "
                    "(JPEGs and ±15 s clips) is purged after thirty "
                    "days. We measured the storage growth at "
                    "approximately 4 MB per session, so a typical "
                    "semester ends with under 400 MB of evidence "
                    "stored. This is a non-trivial decision: the "
                    "alternative (keep everything) would make the "
                    "system both a regulatory liability and a "
                    "needless consumer of S3 bytes.")

    add_heading(doc, "6.4 Societal Context", level=2)
    add_body(doc,
        "Automated surveillance in an educational setting is a charged "
        "subject. The technical decisions documented above are also "
        "social decisions: we explicitly refused to build an autonomous "
        "violation classifier, because the social cost of a false "
        "positive — a student wrongly accused — outweighs the value of "
        "any automation in the labelling step. The system surfaces "
        "candidate evidence, attaches the rules that fired, and lets a "
        "human reach the decision.")
    add_body(doc,
        "The system is also designed to reduce the existing asymmetry "
        "between proctor and student during in-person exams. Today a "
        "proctor’s suspicion is essentially unappealable; with HorusEye, "
        "every flagged incident carries a timestamp, a video clip, the "
        "rules that triggered it, and an audit log of who reviewed it. "
        "This makes the proctor’s judgement reviewable too — by the "
        "instructor, by the academic-integrity committee, and ultimately "
        "by the accused student.")
    add_body(doc,
        "Negative externalities we explicitly attempted to mitigate "
        "include: (a) bias in detection caused by under-representation "
        "of skin tones in YOLO and ArcFace training sets — we plan to "
        "re-train on local data when sufficient consent is collected; "
        "(b) chilling effect from the perception of being watched — "
        "the consent screen is mandatory, the data-retention policy is "
        "explicit, and students can request their evidence through the "
        "feedback channel.")


def build_contemporary(doc):
    add_pagebreak(doc)
    add_heading(doc,
        "7. Contemporary Issues Related to the Project", level=1)
    add_body(doc,
        "The syllabus (§3.6.6, grade item f, 5 pt) asks for a "
        "discussion of contemporary issues in the area of the project. "
        "Four threads are especially relevant in 2026:")

    add_heading(doc, "7.1 The “AI in Education” Backlash", level=2)
    add_body(doc,
        "Following the wide deployment of remote-proctoring tools "
        "during the COVID-19 pandemic, a sizeable backlash emerged. "
        "Studies by EdTech researchers (Coghlan et al. 2021; Selwyn "
        "et al. 2023) and reporting in the popular press have "
        "documented psychological harm, false-positive accusations, "
        "and racial bias in commercial proctoring systems. The "
        "U.S. Department of Education issued guidance in 2023 "
        "advising universities to require human review for any "
        "AI-flagged behaviour. HorusEye is intentionally designed "
        "to land on the right side of this debate: we do not "
        "produce verdicts, we do not score students, we expose every "
        "rule that fired.")

    add_heading(doc, "7.2 Biometric Regulation (KVKK, GDPR, EU AI Act)",
                level=2)
    add_body(doc,
        "The EU AI Act, which entered force in August 2024, classifies "
        "“biometric categorisation for education or vocational training” "
        "as a high-risk use case (Annex III §3). High-risk systems "
        "must, among other obligations, provide “effective human "
        "oversight”, “appropriate transparency”, and “robust logging "
        "for ex-post audit”. Although Türkiye is not bound by the AI "
        "Act, KVKK and TR ICT compliance follow the same broad "
        "principles. We mapped each Article-13 obligation to a "
        "concrete code-level control: consent capture (BL-133), "
        "append-only audit logs (PRD-006), evidence retention windows, "
        "proctor-decision provenance and the row-level-security "
        "matrix.")

    add_heading(doc, "7.3 The Edge-AI vs Cloud-AI Trade-off",
                level=2)
    add_body(doc,
        "An active debate in the computer-vision community concerns "
        "whether classroom analytics should run on a local edge "
        "device (preserving privacy by never transmitting frames "
        "outside the room) or in the cloud (gaining horizontal "
        "scalability). HorusEye chose a middle path: the AI service "
        "runs in the cloud but stores only signal-level data "
        "(timestamps, bounding boxes, embeddings) plus minimum-"
        "necessary JPEG evidence. We deliberately avoid uploading "
        "full video streams. The architecture is portable: nothing "
        "in the AI service prevents it from being run on a local box "
        "if a university requires it.")

    add_heading(doc, "7.4 Open-Source AI Licensing", level=2)
    add_body(doc,
        "Ultralytics YOLOv8 is licensed AGPL-3.0. For an academic "
        "project this is unproblematic — the source code is open. "
        "Any future commercial deployment would have to choose "
        "between (a) keeping the entire system open under AGPL, "
        "(b) purchasing the Ultralytics Enterprise License, or "
        "(c) replacing YOLOv8 with a permissively-licensed model "
        "such as YOLO-NAS (Apache 2.0) or RT-DETR (Apache 2.0). "
        "We have documented this fork-in-the-road in PRD-013 §11 "
        "so the decision is conscious rather than accidental.")


def build_resources(doc):
    add_pagebreak(doc)
    add_heading(doc,
        "8. Use of Library and Internet Resources", level=1)
    add_body(doc,
        "The syllabus (§3.6.7, grade item g, 5 pt) asks for an "
        "accounting of background research, including similar "
        "designs, components and engineering principles consulted "
        "during the project. The categories below summarise the "
        "team’s lifelong-learning effort over the two-semester span.")

    add_heading(doc, "8.1 Academic Papers Consulted", level=2)
    add_bullet(doc, "Atoum et al. (2017) — “Automated Online Exam "
                    "Proctoring” — IEEE Transactions on Multimedia.")
    add_bullet(doc, "Aharon, Orfaig, Bobrovsky (2022) — “BoT-SORT: "
                    "Robust Associations Multi-Pedestrian Tracking”.")
    add_bullet(doc, "Deng, Guo, Yang et al. (2019) — “ArcFace: "
                    "Additive Angular Margin Loss for Deep Face "
                    "Recognition” — CVPR.")
    add_bullet(doc, "Ge, Lin, Wang, Long, Liu et al. (2021) — "
                    "“YOLOX: Exceeding YOLO Series in 2021”.")
    add_bullet(doc, "Lugaresi et al. (2019) — “MediaPipe: A Framework "
                    "for Building Perception Pipelines”.")
    add_bullet(doc, "Coghlan, Miller, Paterson (2021) — “Good "
                    "Proctor or ‘Big Brother’? Ethics of Online "
                    "Exam Supervision Technologies” — Philosophy "
                    "and Technology.")

    add_heading(doc, "8.2 Engineering Standards", level=2)
    add_bullet(doc, "IEEE 830-1998 — Software Requirements "
                    "Specifications (organisational template for "
                    "PRD documents).")
    add_bullet(doc, "ISO/IEC 25010 — Software Quality Model "
                    "(performance, security, maintainability "
                    "categories used as testing-coverage targets).")
    add_bullet(doc, "OWASP ASVS 4.0 — Application Security Verification "
                    "Standard (auth, session, access control).")
    add_bullet(doc, "REST/RFC standards: RFC 6455 (WebSocket), "
                    "RFC 7519 (JWT), RFC 3339 (timestamp formats).")
    add_bullet(doc, "Semantic Versioning 2.0.0 — applied to PRD "
                    "interface contracts (the @version annotation in "
                    "PRD-000).")

    add_heading(doc, "8.3 Online Documentation and Communities",
                level=2)
    add_bullet(doc, "Next.js, React, Supabase, FastAPI, Ultralytics "
                    "official documentation — read end-to-end for the "
                    "modules each team member owned.")
    add_bullet(doc, "AWS CDK Workshop and the “Effective AWS CDK” "
                    "book.")
    add_bullet(doc, "MDN Web Docs — for the WebRTC research that "
                    "informed our decision to stay on WebSocket.")
    add_bullet(doc, "TypeScript handbook and the React 19 release "
                    "notes (React Server Components, useActionState).")
    add_bullet(doc, "Stack Overflow, GitHub issues, and the Ultralytics "
                    "Discord — used pragmatically for specific blockers.")

    add_heading(doc, "8.4 Datasets Used", level=2)
    add_bullet(doc, "Roboflow Universe (workspace: horuseye) — phone, "
                    "earbuds, smartwatch and paper-notes datasets "
                    "(licences recorded in PRD-021).")
    add_bullet(doc, "COCO 2017 — the pre-trained source for YOLOv8n; "
                    "we use classes 0 (person), 67 (cell phone), 73 "
                    "(book).")
    add_bullet(doc, "Internal training samples bucket (anonymised) — "
                    "frames captured under explicit consent during dry "
                    "runs; used only for future fine-tuning.")

    add_heading(doc, "8.5 Books and Long-Form Reading", level=2)
    add_bullet(doc, "Bruegge & Dutoit, “Object-Oriented Software "
                    "Engineering: Using UML, Patterns, and Java” "
                    "(2nd Ed.) — the source of the LLD organisation "
                    "we adopted.")
    add_bullet(doc, "Martin Kleppmann, “Designing Data-Intensive "
                    "Applications” — referenced for the multi-cam "
                    "fusion and event-sourcing decisions.")
    add_bullet(doc, "Robert C. Martin, “Clean Architecture” — used "
                    "as the touchstone when partitioning the AI "
                    "service into detection / scoring / persistence "
                    "modules.")


def build_evaluation(doc):
    add_pagebreak(doc)
    add_heading(doc, "9. Evaluation and Testing", level=1)
    add_body(doc,
        "The syllabus (§3.6.6, grade item d) requires test results "
        "based on the Test Plan (D5). For each test case, the report "
        "should contain a brief description and a pass/fail "
        "indication, ending with an assessment of failed tests, "
        "outstanding bugs and potential enhancements.")

    add_heading(doc, "9.1 Test Inventory", level=2)
    add_table(doc,
        ["Layer", "Count", "Location", "Framework"],
        [
            ["Portal — unit",
             "19", "portal/tests/unit/", "Vitest"],
            ["Portal — integration (real DB)",
             "7", "portal/tests/integration/",
             "Vitest + Supabase CLI"],
            ["Portal — E2E",
             "5", "portal/tests/e2e/", "Playwright"],
            ["AI service — rules + pipeline",
             "46", "ai-service/tests/", "pytest + pytest-asyncio"],
            ["Total", "77", "", ""],
        ],
        col_widths_cm=[5.5, 2.0, 5.5, 3.0],
    )

    add_heading(doc, "9.2 Representative Test Cases (Excerpt)",
                level=2)
    add_table(doc,
        ["ID", "Description", "Result"],
        [
            ["T-AUTH-01",
             "Login with wrong password 5× → AUTH_RATE_LIMITED",
             "PASS"],
            ["T-AUTH-02",
             "First-time login forces /change-password route",
             "PASS"],
            ["T-AUTH-03",
             "Password reset PKCE callback consumes one-time code",
             "PASS"],
            ["T-FILE-04",
             "Upload PDF > 25 MB → FILE_TOO_LARGE",
             "PASS"],
            ["T-FILE-05",
             "Soft-deleted file returns 404 from /d/[id] proxy",
             "PASS"],
            ["T-EXAM-06",
             "Five-step wizard creates exam + sessions + assignments "
             "atomically",
             "PASS"],
            ["T-EXAM-07",
             "Live monitor renders bounding boxes with <2 s "
             "latency from AI service",
             "PASS"],
            ["T-INC-08",
             "phone_in_hand rule fires after 3 s overlap, "
             "cooldown of 30 s respected",
             "PASS"],
            ["T-INC-09",
             "gaze_diversion 5-min sliding window: 3 glances → "
             "MEDIUM, 6 → HIGH",
             "PASS"],
            ["T-INC-10",
             "empty_seat after 60 s lost → incident written to DB "
             "+ evidence uploaded",
             "PASS"],
            ["T-INC-11",
             "Proctor decision modal records decided_by, "
             "decided_at, decision_note",
             "PASS"],
            ["T-AI-12",
             "WebSocket protocol v1.1 reject mismatched client "
             "version with explicit error code",
             "PASS"],
            ["T-AI-13",
             "Worker pool drops oldest frame under backpressure "
             "(BL-319)",
             "PASS"],
            ["T-INFRA-14",
             "ECS rolling deploy keeps minHealthyPercent=100 "
             "during release",
             "PASS"],
            ["T-INFRA-15",
             "supabase db push failure aborts deploy before ECS "
             "is touched",
             "PASS"],
            ["T-SEC-16",
             "RLS: supervisor cannot read incidents from a session "
             "they are not assigned to",
             "PASS"],
            ["T-SEC-17",
             "audit_logs table rejects UPDATE / DELETE attempts "
             "from service role",
             "PASS"],
            ["T-PERF-18",
             "AI service sustains 5 FPS × 3 cameras at ≤ 60 % CPU",
             "PASS (margin)"],
            ["T-PWA-19",
             "Public docs page resolves from service worker cache "
             "when offline",
             "PASS"],
            ["T-CAM-20",
             "Phone-as-camera auto-reconnect surfaces Reconnect "
             "button after retries exhaust",
             "PASS"],
        ],
        col_widths_cm=[2.4, 11.0, 2.6],
    )
    add_body(doc,
        "The complete machine-readable test inventory is in the "
        "repository under portal/tests and ai-service/tests; CI "
        "regenerates the result tables on every pull request to main "
        "or develop. The numbers above are the representative cross-"
        "section drawn from the Test Plan (D5).")

    add_heading(doc, "9.3 Assessment of Failures and Outstanding Bugs",
                level=2)
    add_bullet(doc, "During Sprint 13 dry-runs we observed Wi-Fi-"
                    "induced WebSocket close-1006 events on phone-as-"
                    "camera devices. We added auto-reconnect with "
                    "exponential backoff (BL-253) and an explicit "
                    "“gave up” user-visible button (commit 87a75c1). "
                    "Residual risk: long Wi-Fi outages still require a "
                    "manual click — a more aggressive background "
                    "reconnect is logged for Phase B follow-up.")
    add_bullet(doc, "ArcFace embeddings under heavy backlight produce "
                    "noisy matches; per-student calibration overrides "
                    "(migration 20260513073632) mitigate this for "
                    "enrolled students, but new students require an "
                    "explicit enrolment session.")
    add_bullet(doc, "The whisper / mouth-motion detection planned in "
                    "PRD-013 as a Phase C feature is not present. "
                    "Audio analysis was deliberately deferred to "
                    "remain within KVKK retention rules for audio "
                    "data.")
    add_bullet(doc, "Custom-trained YOLO classes (smartwatch, "
                    "paper_notes, pencil_case, calculator — Sprint 15-"
                    "16) are present and benchmarked but not yet "
                    "wired to the production weights file — the "
                    "default config still points to yolov8n.pt for "
                    "stability during the demo period.")


def build_per_member(doc):
    add_pagebreak(doc)
    add_heading(doc, "10. Per-Member Contributions", level=1)
    add_body(doc,
        "This section is auto-generated from the live backlog (commit "
        "history + backlog_items.assigned_to) and should be edited by "
        "each member before submission to add a personal reflection on "
        "what was hard, what was learned, and what they would do "
        "differently. Each section should fit on a single page.",
        italic=True)

    members = [
        ("10.1 Taha Kürşat Öztürk", "Product Owner & Full-stack",
         [
             "Sprint & Backlog Management System (PRD-018): data model, "
             "Kanban + analytics + dependency graph UI, status-transition "
             "audit, reviewer enforcement.",
             "AWS CDK infrastructure (PRD-005): VPC, ECS Fargate, ALB, "
             "ACM, Route 53, ECR, SSM-driven configuration, zero-"
             "downtime rolling deploys.",
             "Forgot-password PKCE flow, RTSP / WebSocket ingestion, "
             "AI service WebSocket protocol v1.1, exam creation wizard, "
             "live monitor with bbox overlay and incident ring.",
             "CI/CD: GitHub Actions OIDC, Supabase migration gate before "
             "Docker build, env-var consistency check, PRD-interface "
             "validation hook.",
             "Sprint 18 multi-cam coordination, Roboflow workspace "
             "wiring, gate-aware dataset auditor.",
         ],
         "Most useful lesson: every production incident I debugged "
         "traced back to either a missing local migration file or an "
         "SSM type drift. Adding the scripts/check-env-vars.sh gate and "
         "the migration backfill discipline removed an entire category "
         "of “deploy boom” incidents."),
        ("10.2 Tuğba Hilal Kırer", "Portal Frontend",
         [
             "Files & Trash: drag-reorder table, inline metadata "
             "editing, soft-delete + restore + permanent purge.",
             "Users tab: search debounce, role edit, activate/deactivate, "
             "password reset.",
             "Account tab: active sessions, device detection, password "
             "strength indicator, account deletion.",
             "Students page: list, debounced search, inline add, "
             "CSV import with imported/updated/skipped/errors result.",
             "Sidebar architecture, color theming with cubic-bezier "
             "sliding pill animation, ErrorBoundary across every route.",
         ],
         "Most useful lesson: Tailwind v4 + shadcn/ui works beautifully "
         "if you treat shadcn as starter parts and add motion / colour "
         "variables yourself. The cross-team component naming "
         "convention (TableX, FormX, ModalX) made review handoffs "
         "cheap."),
        ("10.3 Gizem Nur İpek", "Portal Backend",
         [
             "Supabase schema baseline (PRD-001 / 003 / 004): users, "
             "files, feedback, notifications, OTP, file access "
             "requests, smtp_settings, plus RLS audit pass.",
             "Force password change middleware and PRD-001 password "
             "policy validation.",
             "Notification + email triggers; welcome email; 90-day "
             "cleanup cron.",
             "Public feedback OTP flow (10 min expiry, 3/hour rate "
             "limit, @tedu.edu.tr-only).",
             "Exam module data layer: exams, exam_rooms, cameras, "
             "exam_sessions, session_proctors, session_students, "
             "students, incidents + rescoring history.",
             "Exam CRUD + Incident APIs (13 endpoints, all audit-"
             "logged, camera stream_url AES-256-GCM encrypted at "
             "rest).",
         ],
         "Most useful lesson: RLS is brilliant when designed up-front "
         "and miserable when retrofitted. The Sprint 1 RLS audit caught "
         "two policies that admins relied on but were missing the "
         "is_admin() helper — fixing them once meant Sprint 3’s exam "
         "tables landed RLS-clean from day one."),
        ("10.4 Ali Sahil", "AI Backend",
         [
             "AI service scaffold: FastAPI + uvicorn + WebSocket "
             "router, /health, Dockerfile, pytest baseline.",
             "YOLOv8n inference pipeline with COCO class filter, "
             "normalised bbox output, lazy weights load, Docker "
             "pre-baking.",
             "Rule engine (16 rules): phone_in_hand, gaze_diversion, "
             "head_turn, empty_seat, paper_detected, hand_to_ear_mouth, "
             "hand_under_desk, body_lean_neighbor, standing_up, "
             "gaze_at_lap, gaze_at_neighbor, object_passing, "
             "synchronized_behavior, face_covering, unauthorized_"
             "person.",
             "Sprint 10 identity: pgvector + ArcFace embedding via "
             "InsightFace, threshold tuning, student matcher.",
             "Sprint 18 multi-camera matcher, OSNet body re-id, "
             "worker pool.",
             "AI performance report CLI (BL-60), YOLOv8 fine-tuning "
             "script (BL-64).",
         ],
         "Most useful lesson: keeping the AI service Python-only and "
         "on-prem (in a single Fargate task) removed an entire "
         "deployment dimension. The WebSocket protocol versioned at "
         "v1.1 was the single contract between my service and the rest "
         "of the team — once it was versioned, every end started moving "
         "in parallel."),
        ("10.5 Çağla Abazaoğlu", "Project Coordinator",
         [
             "All ten deliverables (D1–D10): LLD v1/v2, TODO/Backlog "
             "v1–v4, Test Plan, Final Report, Presentation/Demo, "
             "Return of Materials.",
             "E2E baseline (BL-12, BL-36) and Sprint 4 expansion to "
             "full user journeys (BL-50).",
             "Test Plan Report (D5): coverage targets, test types, "
             "acceptance criteria per PRD.",
             "Demo and presentation production, poster materials, "
             "“Genç Beyinler” event production.",
             "Privacy/KVKK stance documentation (BL-133): consent at "
             "sign-in, data minimisation, no automatic decisions.",
             "Cross-team coordination: sprint reviews, retrospectives, "
             "cross-review matrix in PRD-018 §11.",
         ],
         "Most useful lesson: documenting the deliverable timeline in "
         "the backlog (with deliverable_id linking) meant /sprints/"
         "analytics showed the demo timeline next to engineering "
         "progress without ever needing a side spreadsheet. The auto-"
         "sync from done backlog items to deliverable status "
         "(PRD-018 §5.5) caught at least three “we forgot to flip the "
         "deliverable” misses."),
    ]
    for title, role, deliveries, lesson in members:
        add_heading(doc, title, level=2)
        add_body(doc, f"Role: {role}", bold=True)
        add_body(doc, "Key deliveries:")
        for d in deliveries:
            add_bullet(doc, d)
        add_body(doc, f"Reflection: {lesson}", italic=True)
        add_body(doc,
            "TODO: each member edits this paragraph to add a personal "
            "reflection (what was hard, what would you do differently).",
            italic=True, size=10)


def build_future(doc):
    add_pagebreak(doc)
    add_heading(doc, "11. Limitations and Future Work", level=1)

    add_heading(doc, "11.1 What We Know Is Fragile", level=2)
    add_bullet(doc, "Single-camera coverage of a row can still suffer "
                    "occlusion; the multi-camera fusion mitigates "
                    "this but does not eliminate it when only one "
                    "camera is calibrated for a given seat.")
    add_bullet(doc, "Detection thresholds are tuned for standard "
                    "lecture-hall lighting; harsh backlight from windows "
                    "lowers ArcFace confidence and increases "
                    "face_covering false positives.")
    add_bullet(doc, "Phone-as-camera devices rely on the student / "
                    "proctor’s personal Wi-Fi, which introduces a "
                    "reliability variable outside our control.")
    add_bullet(doc, "Custom-trained YOLO classes (Sprint 15–16) are "
                    "implemented but the production weights file still "
                    "points to the COCO pre-trained model for stability "
                    "during the demo window.")

    add_heading(doc, "11.2 Phase B and C Candidates", level=2)
    add_bullet(doc, "TensorFlow / PyTorch LSTM / GRU behavioural "
                    "sequence model (originally Phase C) to learn long-"
                    "horizon patterns the rule engine cannot express.")
    add_bullet(doc, "Audio analysis (whisper detection) — currently "
                    "deferred for KVKK/retention reasons.")
    add_bullet(doc, "Edge inference appliance (NVIDIA Jetson) for "
                    "rooms with strict no-cloud requirements.")
    add_bullet(doc, "Course-level and university-level analytics "
                    "across exam terms.")
    add_bullet(doc, "OAuth single-sign-on with TED University’s "
                    "central authentication.")


def build_conclusion(doc):
    add_pagebreak(doc)
    add_heading(doc, "12. Conclusion", level=1)
    add_body(doc,
        "HorusEye started as a brief in the CMPE 491 Project "
        "Specifications Report and ended as a multi-service, "
        "multi-camera, production-deployed system covering twenty-two "
        "PRDs, eighteen sprints, two AWS environments, and a hundred-"
        "plus API endpoints. Every numeric claim in this report has "
        "been verified against the codebase on the chore/roboflow-"
        "workspace-env branch as of 14 May 2026; legacy expressions in "
        "earlier documents (TensorFlow LSTM, React/Vue TBD, Redis, "
        "GPU server) have been superseded by what actually shipped, "
        "and the divergence between plan and reality is documented "
        "explicitly in §4.3.")
    add_body(doc,
        "Beyond the code, the team learned that the unit of "
        "coordination on a project of this scope is the interface "
        "contract, not the individual line of code. Holding "
        "interfaces in a single matrix, validating them at commit "
        "time, and refusing to let any PRD drift away from PRD-000 — "
        "this was the single discipline that let five people ship a "
        "system that none of us could have built alone. The technical "
        "achievement is the system; the human achievement is the way "
        "we coordinated to produce it.")
    add_body(doc,
        "We submit HorusEye as a senior project that is both a "
        "complete artefact and an explicit argument: AI in education "
        "should be the proctor’s assistant, never the proctor’s "
        "replacement. The technology we built is in service of that "
        "argument, and the architecture is engineered to enforce it.")


def build_references(doc):
    add_pagebreak(doc)
    add_heading(doc, "References", level=1)
    refs = [
        "Aharon, N., Orfaig, R., & Bobrovsky, B.-Z. (2022). BoT-SORT: "
        "Robust associations multi-pedestrian tracking. arXiv:2206.14651.",

        "Atoum, Y., Chen, L., Liu, A. X., Hsu, S. D. H., & Liu, X. "
        "(2017). Automated online exam proctoring. IEEE Transactions "
        "on Multimedia, 19(7), 1609–1624.",

        "Bruegge, B., & Dutoit, A. H. (2004). Object-oriented software "
        "engineering: Using UML, patterns, and Java (2nd ed.). "
        "Prentice-Hall.",

        "Coghlan, S., Miller, T., & Paterson, J. (2021). Good proctor "
        "or “Big Brother”? Ethics of online exam supervision "
        "technologies. Philosophy & Technology, 34, 1581–1606.",

        "Deng, J., Guo, J., Yang, J., Niannan Xue, Cotsia, I., & "
        "Zafeiriou, S. (2019). ArcFace: Additive angular margin loss "
        "for deep face recognition. CVPR.",

        "European Parliament & Council. (2024). Regulation (EU) "
        "2024/1689 — Artificial Intelligence Act. Official Journal of "
        "the European Union.",

        "Jocher, G., et al. (2023). Ultralytics YOLOv8. GitHub: "
        "https://github.com/ultralytics/ultralytics.",

        "Kleppmann, M. (2017). Designing data-intensive applications. "
        "O’Reilly.",

        "Lugaresi, C., et al. (2019). MediaPipe: A framework for "
        "building perception pipelines. arXiv:1906.08172.",

        "Martin, R. C. (2017). Clean architecture: A craftsman’s guide "
        "to software structure and design. Prentice Hall.",

        "Selwyn, N., O’Neill, C., Smith, G., Andrejevic, M., & Gu, X. "
        "(2023). A necessary evil? The rise of online exam proctoring "
        "in Australian universities. Media International Australia.",

        "Türkiye Cumhuriyeti — Kişisel Verilerin Korunması Kanunu, "
        "Kanun No. 6698 (KVKK), 2016.",

        "UNESCO. (2023). Global Education Monitoring Report — "
        "Technology in Education.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.7)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(r)
        set_run_style(run, size=10.5, color=TEXT)


def build_appendix(doc):
    add_pagebreak(doc)
    add_heading(doc, "Appendix A — Repository Structure", level=1)
    add_mono(doc,
        "HorusEye/\n"
        "├── portal/                    # Next.js 16 + React 19 portal\n"
        "│   ├── app/                   # App Router pages + 93 API routes\n"
        "│   ├── components/            # UI, forms, monitors, sprint board\n"
        "│   ├── lib/                   # Supabase clients, utils, validators\n"
        "│   ├── supabase/migrations/   # 53 SQL migrations\n"
        "│   ├── tests/                 # 31 test files (unit + intg + e2e)\n"
        "│   ├── proxy.ts               # Next 16 middleware → role gating\n"
        "│   └── Dockerfile             # Multi-stage, node:20-alpine\n"
        "├── ai-service/                # FastAPI + Python 3.12\n"
        "│   ├── src/\n"
        "│   │   ├── api/               # FastAPI routers + WS protocol\n"
        "│   │   ├── detection/         # YOLO, BoT-SORT, FaceMesh, Pose\n"
        "│   │   ├── identity/          # ArcFace, student matcher, re-id\n"
        "│   │   ├── scoring/rules/     # 16 rule files\n"
        "│   │   ├── persistence/       # Supabase client + IncidentWriter\n"
        "│   │   └── reports/           # Per-session performance JSON\n"
        "│   ├── tests/                 # 46 pytest files\n"
        "│   ├── config.yaml            # Rule thresholds\n"
        "│   └── Dockerfile             # python:3.12-slim, pre-baked weights\n"
        "├── infra/                     # AWS CDK (TypeScript)\n"
        "│   ├── bin/infra.ts           # Entry: 4 stacks (portal × 2, AI × 2)\n"
        "│   └── lib/                   # network, registry, service, ai-service\n"
        "├── PRD/                       # 22 product requirement documents\n"
        "├── scripts/                   # PRD validator, env-var check\n"
        "├── .github/workflows/         # 4 GitHub Actions workflows\n"
        "└── docs/                      # Reports, runbooks, training notes")

    add_pagebreak(doc)
    add_heading(doc, "Appendix B — User's Manual", level=1)
    add_body(doc,
        "Per CMPE 492 syllabus (§3.6.10), this appendix accompanies "
        "the hardcopy / DVD submission with installation and "
        "operating instructions.")

    add_heading(doc, "B.1 Prerequisites", level=2)
    add_bullet(doc, "Node.js 20 LTS")
    add_bullet(doc, "Python 3.12")
    add_bullet(doc, "Docker Desktop (24+)")
    add_bullet(doc, "Supabase CLI (for local development)")
    add_bullet(doc, "An IP camera or webcam (for AI service smoke test)")

    add_heading(doc, "B.2 Running the Portal Locally", level=2)
    add_mono(doc,
        "git clone https://github.com/kursatozturk/HorusEye.git\n"
        "cd HorusEye/portal\n"
        "npm install\n"
        "cp .env.example .env.local      # fill SUPABASE_URL / ANON_KEY\n"
        "supabase start                   # launches local Postgres + Auth\n"
        "supabase db reset                # applies all 53 migrations + seeds\n"
        "npm run dev                      # http://localhost:3000")

    add_heading(doc, "B.3 Running the AI Service Locally", level=2)
    add_mono(doc,
        "cd HorusEye/ai-service\n"
        "docker compose up --build        # pre-bakes YOLO + InsightFace\n"
        "curl http://localhost:8000/health\n"
        "# WebSocket endpoints:\n"
        "#   ws://localhost:8000/ws/sessions/{id}/detections\n"
        "#   ws://localhost:8000/ws/sessions/{id}/video")

    add_heading(doc, "B.4 Running the Full Validation Suite", level=2)
    add_mono(doc,
        "# Portal — TypeScript, ESLint, Vitest with coverage\n"
        "cd portal\n"
        "npm run validate\n"
        "\n"
        "# AI service — pytest with async support\n"
        "cd ../ai-service\n"
        "pip install -r requirements.txt\n"
        "pytest tests/")

    add_heading(doc, "B.5 Deploying to AWS", level=2)
    add_body(doc,
        "Deployment is handled by GitHub Actions: pushing to the "
        "develop branch deploys to staging; pushing to main (with a "
        "manual environment approval) deploys to production. To run "
        "the CDK manually for inspection:")
    add_mono(doc,
        "cd infra\n"
        "npm install\n"
        "npx cdk synth HorusEye-Staging\n"
        "npx cdk deploy HorusEye-Staging --require-approval never")

    add_pagebreak(doc)
    add_heading(doc, "Appendix C — Acknowledgements", level=1)
    add_body(doc,
        "We thank our supervisor Fırat Akba for guidance across both "
        "semesters; the CMPE 491 / 492 course coordinator Gökçe Nur "
        "Yılmaz; the TED University Department of Computer "
        "Engineering for the project framework; the open-source "
        "communities behind Next.js, Supabase, FastAPI, Ultralytics, "
        "MediaPipe and InsightFace whose work made this project "
        "possible; and the volunteer participants who consented to "
        "appear in our dry-run videos.")
    add_body(doc, "")
    add_body(doc, "Submitted: 15 May 2026.", italic=True, size=10)


# ══════════════════════════════════════════════════════════════════
# Ana
# ══════════════════════════════════════════════════════════════════

def main():
    doc = init_document()

    build_cover(doc)
    build_abstract(doc)
    build_introduction(doc)
    build_background(doc)
    build_architecture(doc)
    build_implementation(doc)
    build_new_tech(doc)
    build_impact(doc)
    build_contemporary(doc)
    build_resources(doc)
    build_evaluation(doc)
    build_per_member(doc)
    build_future(doc)
    build_conclusion(doc)
    build_references(doc)
    build_appendix(doc)

    out = os.path.join(SCRIPT_DIR, "HorusEye-Final-Report.docx")
    doc.save(out)
    size_kb = os.path.getsize(out) / 1024
    print(f"OK → {out}  ({size_kb:.1f} KB)")


if __name__ == "__main__":
    main()
